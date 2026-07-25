#!/usr/bin/env python3
"""
Resume Generator for Lead Data Engineer at PepsiCo
Creates a professional 2-page resume in .docx format
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def add_heading(doc, text, level=1):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_section_title(doc, title):
    """Add a section title with underline"""
    heading = doc.add_heading(title, level=2)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(6)
    heading_format.space_after = Pt(3)
    return heading


def create_resume():
    """Create professional resume document"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run("LEAD DATA ENGINEER")
    name_run.font.size = Pt(16)
    name_run.font.bold = True

    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.text = "Data-Driven Professional | PepsiCo | Cloud Solutions Architect\nEmail: your.email@pepsico.com | Phone: (XXX) XXX-XXXX | LinkedIn: linkedin.com/in/yourprofile"
    contact_format = contact.paragraph_format
    contact_format.space_after = Pt(6)

    # Professional Summary
    add_section_title(doc, "PROFESSIONAL SUMMARY")
    summary = doc.add_paragraph(
        "Results-driven Lead Data Engineer with 8+ years of experience architecting and deploying "
        "enterprise-scale data solutions. Proven expertise in designing cloud-native data platforms on AWS, Azure, "
        "and Databricks. Skilled in building real-time analytics pipelines, data warehouses, and AI/ML infrastructure. "
        "Track record of reducing infrastructure costs by 40% and improving data processing efficiency by 60%. "
        "Passionate about leveraging data engineering to drive business insights and operational excellence at PepsiCo."
    )
    summary.paragraph_format.space_after = Pt(6)

    # Core Competencies
    add_section_title(doc, "CORE COMPETENCIES")
    competencies = [
        "Cloud Platforms: Azure (Synapse, Data Lake, HDInsight), AWS (S3, EC2, RDS, Glue, EMR, Redshift), Databricks",
        "Data Engineering: Apache Spark, Delta Lake, Apache Kafka, ETL/ELT Pipeline Design, Data Warehousing (Snowflake, Redshift)",
        "Programming: Python, Scala, SQL, PySpark, Java",
        "Big Data Technologies: Hadoop, HDFS, Hive, PySpark, Spark Streaming",
        "Tools & Platforms: Azure Data Factory, AWS Glue, Airflow, dbt, Power BI, Tableau",
        "Databases: PostgreSQL, MongoDB, Cosmos DB, DynamoDB, SQL Server",
        "DevOps & CI/CD: Docker, Kubernetes, Git, Jenkins, Azure DevOps, GitLab CI",
        "Data Governance: Data Quality, Lineage Tracking, Metadata Management, Compliance (GDPR, CCPA)"
    ]

    for comp in competencies:
        p = doc.add_paragraph(comp, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    # Professional Experience
    add_section_title(doc, "PROFESSIONAL EXPERIENCE")

    # Job 1
    job1 = doc.add_paragraph()
    job1_title = job1.add_run("Lead Data Engineer, PepsiCo Global")
    job1_title.bold = True
    job1.add_run(" | New York, NY | 2020 – Present")
    job1.paragraph_format.space_after = Pt(2)

    achievements1 = [
        "Architected and deployed a multi-cloud data platform on Azure Synapse and AWS S3/Redshift, reducing query latency by 65%",
        "Led team of 5 data engineers in designing and implementing real-time streaming pipelines using Kafka and Spark Streaming for 50+ data sources",
        "Implemented Databricks lakehouse architecture, enabling self-service analytics for 200+ business users across organization",
        "Designed data governance framework using Apache Atlas and custom metadata management, ensuring 99.9% data quality SLA",
        "Optimized Spark jobs and Delta tables, reducing compute costs by $2.3M annually and improving performance by 70%",
        "Mentored junior engineers and established data engineering best practices, improving deployment frequency by 300%"
    ]

    for achievement in achievements1:
        p = doc.add_paragraph(achievement, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    # Job 2
    job2 = doc.add_paragraph()
    job2_title = job2.add_run("Senior Data Engineer, Fortune 500 Technology Company")
    job2_title.bold = True
    job2.add_run(" | San Francisco, CA | 2017 – 2020")
    job2.paragraph_format.space_after = Pt(2)

    achievements2 = [
        "Designed and implemented enterprise ETL pipelines processing 5TB+ of data daily using AWS Glue and Lambda",
        "Built real-time data warehouse on AWS Redshift supporting 500+ concurrent users with <2s query response times",
        "Established CI/CD pipelines using Jenkins and GitLab CI, reducing deployment time from 4 hours to 15 minutes",
        "Implemented data quality framework reducing data inconsistencies by 85% and improving compliance reporting",
        "Led migration of legacy on-premises data infrastructure to cloud, saving $1.5M in annual operational costs"
    ]

    for achievement in achievements2:
        p = doc.add_paragraph(achievement, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    # Job 3
    job3 = doc.add_paragraph()
    job3_title = job3.add_run("Data Engineer, Software Development Company")
    job3_title.bold = True
    job3.add_run(" | Seattle, WA | 2015 – 2017")
    job3.paragraph_format.space_after = Pt(2)

    achievements3 = [
        "Developed scalable data pipelines using Python and PySpark for processing customer behavioral data",
        "Optimized Hive queries and Hadoop MapReduce jobs, improving batch processing performance by 50%",
        "Designed data models for business intelligence team, enabling creation of 100+ dashboards and reports",
        "Implemented monitoring and alerting system for data pipelines using ELK stack, reducing mean time to detection by 75%"
    ]

    for achievement in achievements3:
        p = doc.add_paragraph(achievement, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    # Certifications
    add_section_title(doc, "CERTIFICATIONS & CREDENTIALS")
    certifications = [
        "Microsoft Certified: Azure Data Engineer Associate (2023)",
        "AWS Certified Solutions Architect – Professional (2023)",
        "Databricks Certified Associate Data Engineer (2022)",
        "AWS Certified Data Analytics – Specialty (2022)",
        "Microsoft Certified: Azure Solutions Architect Expert (2021)",
        "Cloudera Certified Associate Data Analyst (2020)",
        "Apache Spark Developer Certification (2019)",
        "AWS Certified Solutions Architect – Associate (2018)"
    ]

    for cert in certifications:
        p = doc.add_paragraph(cert, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    # Education
    add_section_title(doc, "EDUCATION")
    edu = doc.add_paragraph()
    edu_title = edu.add_run("Master of Science in Computer Science")
    edu_title.bold = True
    edu.add_run(" | University of Washington | Seattle, WA | 2014")
    edu.paragraph_format.space_after = Pt(2)

    edu2 = doc.add_paragraph()
    edu2_title = edu2.add_run("Bachelor of Science in Information Technology")
    edu2_title.bold = True
    edu2.add_run(" | State University | 2012")
    edu2.paragraph_format.space_after = Pt(6)

    # Additional Skills
    add_section_title(doc, "ADDITIONAL SKILLS")
    additional = [
        "Languages: English (Native), Spanish (Conversational)",
        "Soft Skills: Leadership, Team Management, Stakeholder Communication, Agile/Scrum, Technical Documentation",
        "Industry Knowledge: Consumer Goods, Supply Chain Analytics, Real-time Analytics, Data Monetization"
    ]

    for skill in additional:
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    return doc


def main():
    """Generate and save resume"""
    doc = create_resume()

    # Create output directory
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{output_dir}/Lead_Data_Engineer_Resume_{timestamp}.docx"

    # Save document
    doc.save(filename)

    print(f"✓ Resume created successfully!")
    print(f"✓ File saved to: {filename}")
    print(f"✓ Document is 2 pages optimized with professional formatting")
    print(f"✓ Includes: Azure, AWS, Databricks expertise")
    print(f"✓ Includes: 8 professional certifications")
    print(f"✓ Ready to customize with your actual details")

    return filename


if __name__ == "__main__":
    main()
