#!/usr/bin/env python3
"""
Document Preservation Helper
Handles reading, preserving, and extracting data from .doc and .docx files
"""

import os
import json
import hashlib
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional, List


class DocPreserver:
    """Preserve and manage Word documents"""

    def __init__(self, output_base: str = "output"):
        self.output_base = output_base
        self.current_date = datetime.now().strftime("%Y-%m-%d")
        self.output_dir = Path(output_base) / self.current_date
        self.backup_dir = self.output_dir / "doc_backups"
        self._ensure_directories()

    def _ensure_directories(self):
        """Create necessary output directories"""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.backup_dir.mkdir(parents=True, exist_ok=True)

    def calculate_checksum(self, file_path: str) -> str:
        """Calculate SHA256 checksum of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def extract_metadata(self, file_path: str) -> Dict:
        """Extract metadata from Word document"""
        try:
            from docx import Document
        except ImportError:
            return {"error": "python-docx not installed. Install with: pip install python-docx"}

        try:
            doc = Document(file_path)
            props = doc.core_properties

            metadata = {
                "filename": os.path.basename(file_path),
                "file_path": str(file_path),
                "file_size_bytes": os.path.getsize(file_path),
                "author": props.author or "Unknown",
                "created": str(props.created) if props.created else "Unknown",
                "modified": str(props.modified) if props.modified else "Unknown",
                "title": props.title or "Untitled",
                "subject": props.subject or "",
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
            }
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract metadata: {str(e)}"}

    def extract_content(self, file_path: str) -> Dict:
        """Extract text content from Word document"""
        try:
            from docx import Document
        except ImportError:
            return {"error": "python-docx not installed"}

        try:
            doc = Document(file_path)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)

            content = {
                "text_content": "\n".join(paragraphs),
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "tables": tables_data,
                "table_count": len(tables_data),
                "summary": {
                    "word_count": sum(len(p.split()) for p in paragraphs),
                    "character_count": sum(len(p) for p in paragraphs)
                }
            }
            return content
        except Exception as e:
            return {"error": f"Failed to extract content: {str(e)}"}

    def create_backup(self, file_path: str) -> Dict:
        """Create timestamped backup of document"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"error": f"File not found: {file_path}"}

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            stem = file_path.stem
            suffix = file_path.suffix
            backup_filename = f"{stem}_{timestamp}{suffix}"
            backup_path = self.backup_dir / backup_filename

            shutil.copy2(file_path, backup_path)

            return {
                "original_file": str(file_path),
                "backup_path": str(backup_path),
                "backup_filename": backup_filename,
                "timestamp": timestamp,
                "size_bytes": backup_path.stat().st_size,
                "created": datetime.now().isoformat()
            }
        except Exception as e:
            return {"error": f"Failed to create backup: {str(e)}"}

    def preserve_document(self, file_path: str) -> Dict:
        """Complete document preservation workflow"""
        file_path = str(file_path)

        if not Path(file_path).exists():
            return {"error": f"File not found: {file_path}"}

        if not Path(file_path).suffix.lower() in [".doc", ".docx"]:
            return {"error": "File must be .doc or .docx format"}

        result = {
            "status": "success",
            "timestamp": datetime.now().isoformat(),
            "metadata": self.extract_metadata(file_path),
            "content_summary": self.extract_content(file_path),
            "backup": self.create_backup(file_path),
            "checksum": self.calculate_checksum(file_path)
        }

        return result

    def save_report(self, preservation_data: Dict, filename: Optional[str] = None) -> str:
        """Save preservation report to JSON file"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"doc_report_{timestamp}.json"

        report_path = self.output_dir / filename

        with open(report_path, "w") as f:
            json.dump(preservation_data, f, indent=2)

        return str(report_path)


def main():
    """Example usage"""
    preserver = DocPreserver()
    print(f"Document Preservation Helper initialized")
    print(f"Output directory: {preserver.output_dir}")


if __name__ == "__main__":
    main()
